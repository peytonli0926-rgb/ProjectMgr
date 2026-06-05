"""领导汇报版 Word 报告生成器。

风格：简洁、结论先行、突出风险等级和整改优先级。
"""

import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..config import RISK_LABELS

logger = logging.getLogger(__name__)

# 颜色
COLOR_HIGH = RGBColor(0xDC, 0x26, 0x26)
COLOR_MEDIUM = RGBColor(0xD9, 0x77, 0x06)
COLOR_LOW = RGBColor(0x15, 0x80, 0x3D)
COLOR_INFO = RGBColor(0x66, 0x70, 0x85)
COLOR_BRAND = RGBColor(0x0D, 0x94, 0x88)
COLOR_HEADING = RGBColor(0x11, 0x1C, 0x34)


def generate_executive_report(
    evidence_data: dict[str, Any],
    source_zip_name: str,
    output_path: str | Path,
) -> Path:
    """生成领导汇报版 Word 报告。"""
    doc = Document()
    out_path = Path(output_path)

    # ── 封面标题 ──
    p = doc.add_heading("Oracle 巡检分析报告（领导决策版）", level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = evidence_data.get("metadata", {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"源文件: {source_zip_name}\n分析时间: {meta.get('analyzed_at', '')}")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_INFO

    doc.add_paragraph()  # 空行

    # ── 总体结论（风险摘要） ──
    doc.add_heading("一、总体结论", level=1)
    risk_summary = evidence_data.get("risk_summary", {})
    total = sum(risk_summary.values())
    p = doc.add_paragraph()
    run = p.add_run(f"本次分析共检测到 {meta.get('total_evidence', 0)} 条证据，覆盖 {meta.get('files_analyzed', 0)} 个文件。")
    run.font.size = Pt(11)

    # 风险等级面板
    doc.add_heading("风险等级分布", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    labels = [("严重", risk_summary.get("critical", 0), COLOR_HIGH),
              ("高", risk_summary.get("high", 0), COLOR_HIGH),
              ("中", risk_summary.get("medium", 0), COLOR_MEDIUM),
              ("低", risk_summary.get("low", 0), COLOR_LOW),
              ("参考", risk_summary.get("info", 0), COLOR_INFO)]
    for i, (label, count, color) in enumerate(labels):
        hdr[i].text = f"{label}: {count}"
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = color

    doc.add_paragraph()

    # ── 各方向问题速览 ──
    doc.add_heading("二、各方向问题速览", level=1)
    by_cat = evidence_data.get("by_category", {})
    for cat, ev_list in by_cat.items():
        high_count = sum(1 for e in ev_list if e.get("severity") in ("critical", "high"))
        medium_count = sum(1 for e in ev_list if e.get("severity") == "medium")
        p = doc.add_paragraph()
        run = p.add_run(f"【{cat}】")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_HEADING
        parts = []
        if high_count:
            parts.append(f"高危 {high_count} 项")
        if medium_count:
            parts.append(f"中危 {medium_count} 项")
        if not parts:
            parts.append("未发现异常")
        run2 = p.add_run(f" — {'，'.join(parts)}")
        run2.font.size = Pt(11)

    doc.add_paragraph()

    # ── 高风险问题详述 ──
    doc.add_heading("三、高风险问题详述", level=1)
    has_high = False
    for ev in evidence_data.get("evidence", []):
        if ev.get("severity") in ("critical", "high"):
            has_high = True
            p = doc.add_paragraph()
            run = p.add_run(f"[{RISK_LABELS.get(ev['severity'], ev['severity'])}] {ev['title']}")
            run.bold = True
            run.font.color.rgb = COLOR_HIGH
            run.font.size = Pt(11)
            p2 = doc.add_paragraph(ev.get("log_snippet", ""), style="List Bullet")
            p2.style.font.size = Pt(9)
            p3 = doc.add_paragraph(f"来源: {ev.get('log_file', '')}  建议: {ev.get('recommendation', '')}")
            p3.runs[0].font.size = Pt(9)
            p3.runs[0].font.color.rgb = COLOR_INFO

    if not has_high:
        doc.add_paragraph("本次分析未发现高风险问题。")

    doc.add_paragraph()

    # ── 整改优先级建议 ──
    doc.add_heading("四、整改优先级建议", level=1)
    recommendations = []
    for ev in evidence_data.get("evidence", []):
        rec = ev.get("recommendation", "").strip()
        if rec and rec not in recommendations:
            recommendations.append(rec)
    for i, rec in enumerate(recommendations[:10], 1):
        p = doc.add_paragraph(f"{i}. {rec}")
        p.style.font.size = Pt(10)

    if not recommendations:
        doc.add_paragraph("本次分析没有需要优先整改的问题。")

    doc.add_paragraph()
    p = doc.add_paragraph("— 报告由 Oracle TFA Analyzer 自动生成，所有结论均可追溯到 evidence.json —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_INFO

    # 保存
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info("领导汇报版报告已生成: %s", out_path)
    return out_path
