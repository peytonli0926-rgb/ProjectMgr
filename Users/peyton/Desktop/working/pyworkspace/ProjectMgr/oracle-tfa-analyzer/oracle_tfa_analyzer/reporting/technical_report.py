"""技术专家版 Word 报告生成器。

风格：详细、包含日志证据、技术解释、建议命令和整改方案。
"""

import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..config import RISK_LABELS

logger = logging.getLogger(__name__)

COLOR_HIGH = RGBColor(0xDC, 0x26, 0x26)
COLOR_MEDIUM = RGBColor(0xD9, 0x77, 0x06)
COLOR_LOW = RGBColor(0x15, 0x80, 0x3D)
COLOR_INFO = RGBColor(0x66, 0x70, 0x85)
COLOR_HEADING = RGBColor(0x11, 0x1C, 0x34)
COLOR_CODE = RGBColor(0x1E, 0x29, 0x3B)


def _add_code_block(doc, text: str):
    """添加等宽代码块段落。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_CODE


def generate_technical_report(
    evidence_data: dict[str, Any],
    source_zip_name: str,
    output_path: str | Path,
) -> Path:
    """生成技术专家版 Word 报告。"""
    doc = Document()
    out_path = Path(output_path)

    # ── 封面 ──
    p = doc.add_heading("Oracle TFA 日志深度分析报告（技术专家版）", level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = evidence_data.get("metadata", {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"源文件: {source_zip_name}\n"
        f"分析时间: {meta.get('analyzed_at', '')}\n"
        f"文件数: {meta.get('files_analyzed', 0)}  |  规则数: {meta.get('rules_applied', 0)}  |  证据数: {meta.get('total_evidence', 0)}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_INFO

    doc.add_paragraph()

    # ── 1. 环境与配置概要 ──
    doc.add_heading("一、分析概要", level=1)
    p = doc.add_paragraph(
        "本报告基于 TFA 收集的日志文件，通过规则引擎自动分析八大方向。"
        "所有结论均来源于日志中的真实证据，不编造日志中未出现的事实。"
        "每个问题均附有日志原文片段、来源文件、技术解释和整改方案。"
    )

    # ── 2. 风险摘要 ──
    doc.add_heading("二、风险摘要", level=1)
    risk_summary = evidence_data.get("risk_summary", {})
    table = doc.add_table(rows=2, cols=5)
    table.style = "Light Grid Accent 1"
    labels = ["严重", "高", "中", "低", "参考"]
    keys = ["critical", "high", "medium", "low", "info"]
    for i, (label, key) in enumerate(zip(labels, keys)):
        table.rows[0].cells[i].text = label
        table.rows[1].cells[i].text = str(risk_summary.get(key, 0))
    doc.add_paragraph()

    # ── 3. 分类详细发现 ──
    doc.add_heading("三、详细发现与证据", level=1)
    by_cat = evidence_data.get("by_category", {})

    for cat, ev_list in by_cat.items():
        doc.add_heading(f"📌 {cat}", level=2)
        high_count = sum(1 for e in ev_list if e.get("severity") in ("critical", "high"))
        p = doc.add_paragraph(f"共 {len(ev_list)} 条证据，其中高风险 {high_count} 条。")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = COLOR_INFO

        # 按严重程度排序
        severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
        ev_list_sorted = sorted(ev_list, key=lambda e: severity_order.get(e.get("severity", "info"), 99))

        for ev in ev_list_sorted:
            sev = ev.get("severity", "info")
            color = {"critical": COLOR_HIGH, "high": COLOR_HIGH, "medium": COLOR_MEDIUM,
                     "low": COLOR_LOW, "info": COLOR_INFO}.get(sev, COLOR_INFO)
            label = RISK_LABELS.get(sev, sev)

            # 标题行
            p = doc.add_paragraph()
            run = p.add_run(f"[{label}] {ev['title']}  (规则: {ev['rule_id']})")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = color

            # 日志证据
            p = doc.add_paragraph()
            run = p.add_run(f"📄 日志文件: {ev['log_file']}")
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_INFO
            if ev.get("line_number"):
                run2 = p.add_run(f"  (行 {ev['line_number']})")
                run2.font.size = Pt(9)
                run2.font.color.rgb = COLOR_INFO

            _add_code_block(doc, ev.get("log_snippet", ""))

            # 技术解释
            if ev.get("detail"):
                p = doc.add_paragraph()
                run = p.add_run(f"💡 技术解释: {ev['detail']}")
                run.font.size = Pt(9)
                run.italic = True

            # 整改建议
            if ev.get("recommendation"):
                p = doc.add_paragraph()
                run = p.add_run(f"🔧 整改建议: {ev['recommendation']}")
                run.font.size = Pt(9)
                run.bold = True

            doc.add_paragraph()

    # ── 4. 附录 ──
    doc.add_heading("四、附录", level=1)
    doc.add_heading("分析命令参考", level=2)
    cmds = [
        "-- 检查 Alert Log 中的 ORA 错误",
        "grep 'ORA-' $ORACLE_BASE/diag/rdbms/*/*/trace/alert_*.log | grep -v 'ORA-0000'",
        "",
        "-- 检查 CRS 资源状态",
        "crsctl stat res -t",
        "",
        "-- 检查 ASM 磁盘组状态",
        "asmcmd lsdg",
        "",
        "-- 检查 Data Guard 同步状态",
        "dgmgrl sys/<password>@<primary> 'show configuration verbose'",
        "",
        "-- 查看 AWR Top SQL",
        "@?/rdbms/admin/awrrpt.sql",
    ]
    for cmd in cmds:
        if cmd:
            _add_code_block(doc, cmd)
        else:
            doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph("— 报告由 Oracle TFA Analyzer 自动生成，基于 evidence.json 中的真实证据 —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_INFO

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info("技术专家版报告已生成: %s", out_path)
    return out_path
