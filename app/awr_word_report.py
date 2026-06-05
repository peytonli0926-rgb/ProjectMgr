"""
awr_word_report — AWR Word 报告生成器（供 Web 端 /oracle/awr-word-report 调用）

专业金融报告风格：深海蓝主色 + 金色强调 + 层次灰排版。
"""

import json
import re
from datetime import datetime
from pathlib import Path


from .config import TMP_OUTPUT

ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = TMP_OUTPUT
CHART_DIR = OUTPUT_DIR / "charts"
AWR_ANALYSIS_MD = OUTPUT_DIR / "awr_analysis_report.md"
AWR_ANALYSIS_DOCX = OUTPUT_DIR / "awr_analysis_report.docx"
AWR_SUMMARY_JSON = OUTPUT_DIR / "awr_summary.json"

# ── 专业金融报告配色 ──
FONT_CN = "Microsoft YaHei"
FONT_MONO = "Consolas"

COLOR_PRIMARY = "1B3A5C"        # 深海蓝 — 封面/主标题
COLOR_ACCENT = "C8A96E"         # 香槟金 — 装饰线/强调
COLOR_TITLE = "1B3A5C"          # H1
COLOR_H2 = "2C5F8A"             # H2
COLOR_H3 = "4A7BA7"             # H3
COLOR_BODY = "2C3E50"           # 正文
COLOR_MUTED = "7F8C8D"          # 辅助文字
COLOR_TABLE_HEAD = "1B3A5C"     # 表头底
COLOR_BORDER = "D5DDE5"         # 表格边框
COLOR_TABLE_ALT = "F0F3F7"      # 斑马纹

COLOR_CONCLUSION_BG = "EBF0F7"
COLOR_CONCLUSION_ACCENT = "1B3A5C"
COLOR_RISK_BG = "FDEDED"
COLOR_RISK_ACCENT = "C0392B"
COLOR_EVIDENCE_BG = "F2F4F4"
COLOR_EVIDENCE_ACCENT = "7F8C8D"
COLOR_ADVICE_BG = "E8F0E8"
COLOR_ADVICE_ACCENT = "27AE60"

SECTION_TITLES = [
    "总体结论",
    "风险等级判断",
    "数据库负载画像",
    "Top Wait Events 分析",
    "Top SQL 分析",
    "主机资源分析",
    "内存与参数建议",
    "问题点清单",
    "整改建议",
    "后续取证清单",
    "领导汇报摘要",
    "专家交付结论",
]
CHINESE_NUMBERS = "一二三四五六七八九十"


def split_table_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def is_table_line(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def rgb(hex_color: str):
    from docx.shared import RGBColor

    value = hex_color.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, size_pt=10.5, bold=False, color=COLOR_BODY, font_name=FONT_CN):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = rgb(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)


def shade_cell(cell, color: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, right=100, bottom=80, left=100):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("right", right), ("bottom", bottom), ("left", left)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLOR_BORDER):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, size_pt=9, bold=False, color=COLOR_BODY, font_name=FONT_CN):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    add_runs_with_bold(paragraph, str(text), size_pt=size_pt, default_bold=bold, color=color, font_name=font_name)


def add_runs_with_bold(paragraph, text: str, size_pt=10.5, default_bold=False, color=COLOR_BODY, font_name=FONT_CN):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = default_bold
        content = part
        if part.startswith("**") and part.endswith("**"):
            bold = True
            content = part[2:-2]
        run = paragraph.add_run(content)
        set_run_font(run, size_pt=size_pt, bold=bold, color=color, font_name=font_name)


def format_paragraph(paragraph, size_pt=10.5, color=COLOR_BODY):
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt, bold=bool(run.bold), color=color)


def create_document():
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("当前 Python 环境未安装 python-docx，请先安装 python-docx") from exc
    document = Document()
    setup_document_styles(document)
    return document


def setup_document_styles(document):
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.start_type = WD_SECTION_START.NEW_PAGE

    # White page background
    bg = document.element.find(qn("w:background"))
    if bg is None:
        from docx.oxml import OxmlElement
        bg = OxmlElement("w:background")
        document.element.append(bg)
    bg.set(qn("w:color"), "FFFFFF")

    normal = document.styles["Normal"]
    normal.font.name = FONT_CN
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_before = Pt(1)
    normal.paragraph_format.space_after = Pt(4)


def add_page_header_footer(document, basic_info: dict):
    """专业页眉：报告标题 + 金色底部分隔线 / 页脚：页码居中 + 灰线上分隔。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    section = document.sections[0]

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_ACCENT)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    db_name = basic_info.get("DB Name") or "Oracle Database"
    run = hp.add_run(f"Oracle AWR 性能分析报告  |  {db_name}")
    set_run_font(run, size_pt=8.5, bold=False, color=COLOR_MUTED)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = 1  # center

    fp_pr = fp._p.get_or_add_pPr()
    fp_bdr = OxmlElement("w:pBdr")
    ftop = OxmlElement("w:top")
    ftop.set(qn("w:val"), "single")
    ftop.set(qn("w:sz"), "4")
    ftop.set(qn("w:space"), "1")
    ftop.set(qn("w:color"), COLOR_BORDER)
    fp_bdr.append(ftop)
    fp_pr.append(fp_bdr)

    run = fp.add_run("第 ")
    set_run_font(run, size_pt=8.5, color=COLOR_MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    r2 = fp.add_run("PAGE")
    set_run_font(r2, size_pt=8.5, color=COLOR_MUTED)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r2._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r2._r.append(sep)
    r3 = fp.add_run("1")
    set_run_font(r3, size_pt=8.5, color=COLOR_MUTED)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r3._r.append(fld_end)
    r4 = fp.add_run(" 页")
    set_run_font(r4, size_pt=8.5, color=COLOR_MUTED)


def load_awr_basic_info() -> dict:
    if not AWR_SUMMARY_JSON.exists():
        return {}
    try:
        payload = json.loads(AWR_SUMMARY_JSON.read_text(encoding="utf-8"))
        return payload.get("basic_info") or {}
    except Exception:
        return {}


def add_horizontal_line(paragraph, color=COLOR_BORDER, size="6"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    p_pr.append(border)


def add_cover_page(document, basic_info: dict):
    """专业封面：金色装饰线围合 + 数据库信息表（左标签右对齐）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    for _ in range(3):
        document.add_paragraph()

    # Top gold line
    tl = document.add_paragraph()
    tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tl.add_run("━" * 60)
    set_run_font(run, size_pt=8, color=COLOR_ACCENT)

    document.add_paragraph()

    # Main title
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Oracle AWR 性能分析报告")
    set_run_font(run, size_pt=24, bold=True, color=COLOR_PRIMARY)

    document.add_paragraph()

    # Subtitle
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于 AWR / ASH / OS 指标的数据库性能诊断与优化建议")
    set_run_font(run, size_pt=12, color=COLOR_MUTED)

    document.add_paragraph()

    # Bottom gold line
    bl = document.add_paragraph()
    bl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = bl.add_run("━" * 60)
    set_run_font(run, size_pt=8, color=COLOR_ACCENT)

    for _ in range(4):
        document.add_paragraph()

    # DB info table
    info = [
        ("数据库名称", basic_info.get("DB Name") or "详见 AWR 摘要"),
        ("实例名称", basic_info.get("Instance Name") or "详见 AWR 摘要"),
        ("主机名称", basic_info.get("Host Name") or "详见 AWR 摘要"),
        ("分析时间窗口", f"{basic_info.get('Begin Snap') or '详见 AWR 摘要'} — {basic_info.get('End Snap') or '详见 AWR 摘要'}"),
        ("AI 分析模型", basic_info.get("Model Name") or "DeepSeek"),
        ("报告生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]
    table = document.add_table(rows=len(info), cols=2)
    table.autofit = True
    set_table_borders(table, color="FFFFFF")
    for idx, (key, value) in enumerate(info):
        set_cell_text(table.rows[idx].cells[0], key, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
        set_cell_margins(table.rows[idx].cells[0], top=60, bottom=60, left=200, right=60)
        table.rows[idx].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        set_cell_text(table.rows[idx].cells[1], value, size_pt=10.5, color=COLOR_BODY)
        set_cell_margins(table.rows[idx].cells[1], top=60, bottom=60, left=60, right=200)

    document.add_page_break()


def extract_section(markdown: str, title_pattern: str) -> str:
    pattern = re.compile(rf"(?ims)^##\s*(?:\d+[.、]\s*)?{re.escape(title_pattern)}\s*$\n(.*?)(?=^##\s+|\Z)")
    match = pattern.search(markdown)
    return match.group(1).strip() if match else ""


def extract_key_conclusions(markdown: str) -> list[str]:
    section = extract_section(markdown, "总体结论")
    source = section or markdown
    conclusions = []
    for line in source.splitlines():
        stripped = line.strip()
        if stripped.startswith(("- ", "* ")):
            conclusions.append(stripped[2:].strip())
        elif re.match(r"^\d+[.、]\s+", stripped):
            conclusions.append(re.sub(r"^\d+[.、]\s+", "", stripped).strip())
        if len(conclusions) >= 5:
            break
    if not conclusions:
        for line in source.splitlines():
            stripped = line.strip()
            if stripped and not stripped.startswith(("#", "|")):
                conclusions.append(stripped)
            if len(conclusions) >= 5:
                break
    return conclusions or ["Markdown 报告未识别到明确总体结论，请结合 AWR 摘要复核。"]


def extract_first_table_after(markdown: str, title_pattern: str) -> list[str]:
    section = extract_section(markdown, title_pattern)
    lines = section.splitlines()
    for idx, line in enumerate(lines):
        if is_table_line(line) and idx + 1 < len(lines) and is_table_separator(lines[idx + 1]):
            table_lines = [line, lines[idx + 1]]
            for next_line in lines[idx + 2:]:
                if not is_table_line(next_line):
                    break
                table_lines.append(next_line)
            return table_lines
    return []


def add_executive_summary(document, markdown: str):
    add_numbered_heading(document, "结论摘要", 1, force_text="结论摘要")
    add_special_block(document, "核心结论", extract_key_conclusions(markdown), block_type="conclusion")

    risk_table = extract_first_table_after(markdown, "风险等级") or extract_first_table_after(markdown, "风险等级判断")
    if risk_table:
        add_numbered_heading(document, "风险等级", 2, force_text="风险等级")
        add_markdown_table(document, risk_table)
    else:
        add_special_block(document, "风险提示", ["Markdown 报告中未识别到风险等级表，请结合规则引擎结果复核。"], block_type="risk")
    document.add_page_break()


def add_numbered_heading(document, text: str, level: int, force_text: str | None = None):
    """H1: 左侧金色饰线 + 底部灰线; H2/H3: 纯文字。"""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    clean = re.sub(r"^\d+[.、]\s*", "", text).strip()
    if level == 1:
        title = force_text or clean
        paragraph = document.add_paragraph()
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), COLOR_ACCENT)
        p_bdr.append(left)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), COLOR_BORDER)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        run = paragraph.add_run(title)
        set_run_font(run, size_pt=16, bold=True, color=COLOR_TITLE)
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.3
        return

    if level == 2:
        sz, clr = 12, COLOR_H2
    else:
        sz, clr = 11, COLOR_H3
    paragraph = document.add_paragraph()
    run = paragraph.add_run(clean)
    set_run_font(run, size_pt=sz, bold=True, color=clr)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)


def add_markdown_content(document, markdown: str):
    lines = markdown.splitlines()
    section_index = 0
    idx = 0
    while idx < len(lines):
        stripped = lines[idx].rstrip().strip()
        if not stripped:
            idx += 1
            continue
        if is_table_line(stripped) and idx + 1 < len(lines) and is_table_separator(lines[idx + 1].strip()):
            table_lines = [stripped, lines[idx + 1].strip()]
            idx += 2
            while idx < len(lines) and is_table_line(lines[idx].strip()):
                table_lines.append(lines[idx].strip())
                idx += 1
            add_markdown_table(document, table_lines)
            continue
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if title != "Oracle AWR 性能分析报告":
                paragraph = document.add_paragraph()
                paragraph.alignment = 1
                run = paragraph.add_run(title)
                set_run_font(run, size_pt=20, bold=True, color=COLOR_PRIMARY)
        elif stripped.startswith("## "):
            section_index += 1
            title = normalize_heading_text(stripped[3:].strip())
            numbered = f"{CHINESE_NUMBERS[section_index - 1] if section_index <= len(CHINESE_NUMBERS) else section_index}、{title}"
            add_numbered_heading(document, numbered, 1, force_text=numbered)
            insert_chart_if_needed(document, title)
        elif stripped.startswith("### "):
            add_numbered_heading(document, stripped[4:].strip(), 2)
        elif is_special_heading(stripped):
            block_lines, idx = collect_block(lines, idx + 1)
            add_special_block(document, special_title(stripped), block_lines, block_type=special_block_type(stripped))
            continue
        elif stripped.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            add_runs_with_bold(paragraph, stripped[2:].strip(), size_pt=10, default_bold=contains_emphasis_keyword(stripped))
            format_paragraph(paragraph, size_pt=10)
        else:
            paragraph = document.add_paragraph()
            add_runs_with_bold(paragraph, stripped, size_pt=10, default_bold=contains_emphasis_keyword(stripped))
            format_paragraph(paragraph, size_pt=10)
        idx += 1


def normalize_heading_text(text: str) -> str:
    text = re.sub(r"^\d+[.、]\s*", "", text).strip()
    aliases = {"风险等级": "风险等级判断"}
    return aliases.get(text, text)


def contains_emphasis_keyword(text: str) -> bool:
    keywords = ("核心结论", "关键发现", "风险", "建议", "现有 AWR 无法证明", "需要补充")
    return any(keyword in text for keyword in keywords)


def is_special_heading(text: str) -> bool:
    return any(keyword in text for keyword in ("核心结论", "风险提示", "后续需要补充", "专家建议"))


def special_title(text: str) -> str:
    if "核心结论" in text:
        return "核心结论"
    if "风险提示" in text or "风险" in text:
        return "风险提示"
    if "后续" in text or "需要补充" in text:
        return "后续需要补充的证据"
    return "专家建议"


def special_block_type(text: str) -> str:
    if "核心结论" in text:
        return "conclusion"
    if "风险" in text:
        return "risk"
    if "后续" in text or "需要补充" in text:
        return "evidence"
    return "advice"


def collect_block(lines: list[str], start: int) -> tuple[list[str], int]:
    block = []
    idx = start
    while idx < len(lines):
        stripped = lines[idx].strip()
        if stripped.startswith("#") or (is_table_line(stripped) and idx + 1 < len(lines) and is_table_separator(lines[idx + 1].strip())):
            break
        if stripped:
            block.append(re.sub(r"^[-*]\s+", "", stripped))
        idx += 1
    return block, idx


def add_special_block(document, title: str, content, block_type="conclusion"):
    """卡片样式：左侧强调色竖条 + 浅背景。"""
    colors = {
        "conclusion": (COLOR_CONCLUSION_BG, COLOR_CONCLUSION_ACCENT),
        "risk": (COLOR_RISK_BG, COLOR_RISK_ACCENT),
        "evidence": (COLOR_EVIDENCE_BG, COLOR_EVIDENCE_ACCENT),
        "advice": (COLOR_ADVICE_BG, COLOR_ADVICE_ACCENT),
    }
    fill, accent = colors.get(block_type, colors["conclusion"])
    items = content if isinstance(content, list) else [str(content)]
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    set_table_borders(table, color=fill)

    shade_cell(table.rows[0].cells[0], accent)
    set_cell_margins(table.rows[0].cells[0], left=30, right=30)
    table.rows[0].cells[0].width = 1

    shade_cell(table.rows[0].cells[1], fill)
    set_cell_margins(table.rows[0].cells[1], top=100, right=140, bottom=80, left=140)

    paragraph = table.rows[0].cells[1].paragraphs[0]
    run = paragraph.add_run(title)
    set_run_font(run, size_pt=11, bold=True, color=COLOR_PRIMARY)

    for item in items:
        p = table.rows[0].cells[1].add_paragraph(style="List Bullet" if len(items) > 1 else None)
        add_runs_with_bold(p, item, size_pt=10, default_bold=True, color=COLOR_BODY)
        format_paragraph(p, size_pt=10)
    document.add_paragraph()


def add_markdown_table(document, rows: list[str]):
    if len(rows) < 2:
        return
    headers = split_table_row(rows[0])
    data_rows = [split_table_row(row) for row in rows[2:] if is_table_line(row)]
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = True
    table.alignment = 1

    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, size_pt=9, bold=True, color="FFFFFF")
    for data in data_rows:
        row = table.add_row()
        for col in range(len(headers)):
            set_cell_text(row.cells[col], data[col] if col < len(data) else "", size_pt=9)
    style_table(table, headers)
    document.add_paragraph()


def style_table(table, headers: list[str]):
    """专业风格：深色表头 + 斑马纹 + 风险色标。"""
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=60, right=80, bottom=60, left=80)
            if row_idx == 0:
                shade_cell(cell, COLOR_TABLE_HEAD)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size_pt=9, bold=True, color="FFFFFF")
                continue
            if row_idx % 2 == 0:
                shade_cell(cell, COLOR_TABLE_ALT)
            text = cell.text.strip()
            hdr = headers[cell_idx] if cell_idx < len(headers) else ""
            if "风险" in hdr or "等级" in hdr:
                if "高" in text:
                    shade_cell(cell, COLOR_RISK_BG)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run, color=COLOR_RISK_ACCENT, bold=True)
                elif "中" in text:
                    shade_cell(cell, "FEF5E7")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run, color="E67E22", bold=True)
                elif "低" in text:
                    shade_cell(cell, "E8F0E8")
                elif "未" in text or "信息不足" in text:
                    shade_cell(cell, "F2F4F4")
            font_name = FONT_MONO if "SQL ID" in hdr.upper() or "SQL Id" in hdr else FONT_CN
            bld = font_name == FONT_MONO
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size_pt=9, bold=bld or bool(run.bold), color=COLOR_BODY, font_name=font_name)


def insert_chart_if_needed(document, title: str):
    mapping = [
        ("主机资源", "host_cpu.png", "Host CPU 利用率"),
        ("CPU", "host_cpu.png", "Host CPU 利用率"),
        ("负载", "awr_load_profile.png", "Load Profile Per Second"),
        ("等待", "top_wait_events.png", "Top Wait Events 占比"),
        ("Top Wait", "top_wait_events.png", "Top Wait Events 占比"),
        ("Top SQL", "top_sql.png", "Top SQL Elapsed Time"),
        ("SQL", "top_sql.png", "Top SQL Elapsed Time"),
        ("命中率", "instance_efficiency.png", "Instance Efficiency 命中率"),
        ("内存", "instance_efficiency.png", "Instance Efficiency 命中率"),
        ("缓存", "instance_efficiency.png", "Instance Efficiency 命中率"),
        ("Segment", "top_segments.png", "Top Segments 访问热点"),
        ("段", "top_segments.png", "Top Segments 访问热点"),
        ("热点", "top_segments.png", "Top Segments 访问热点"),
    ]
    for keyword, filename, caption in mapping:
        if keyword in title:
            add_chart(document, CHART_DIR / filename, caption)
            return


def add_chart(document, path: Path, caption: str):
    if not path.exists():
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    cap_para = document.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(f"图：{caption}")
    set_run_font(cap_run, size_pt=9, color="666666")


def save_document(document, output_path: Path):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def markdown_to_word(markdown_path: Path = AWR_ANALYSIS_MD, output_path: Path = AWR_ANALYSIS_DOCX, model_name: str | None = None) -> dict:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not markdown_path.exists():
        raise FileNotFoundError("请先执行 AWR 分析，生成 Markdown 报告")

    document = create_document()
    basic_info = load_awr_basic_info()
    if model_name:
        basic_info["Model Name"] = model_name
    add_page_header_footer(document, basic_info)
    add_cover_page(document, basic_info)

    try:
        markdown = markdown_path.read_text(encoding="utf-8")
        add_executive_summary(document, markdown)
        add_markdown_content(document, markdown)
    except Exception as exc:
        add_special_block(
            document, "Word 生成提示",
            [f"Markdown 解析失败，已生成基础 Word。错误信息：{exc}"],
            block_type="risk",
        )
        paragraph = document.add_paragraph()
        paragraph.add_run(markdown_path.read_text(encoding="utf-8", errors="replace"))
        format_paragraph(paragraph)

    save_document(document, output_path)
    return {
        "message": "Word 报告生成成功",
        "markdown_path": str(markdown_path),
        "word_path": str(output_path),
    }
