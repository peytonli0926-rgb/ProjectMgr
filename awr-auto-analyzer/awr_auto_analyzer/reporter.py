"""
awr_auto_analyzer.reporter — 报告生成器

支持两种输出格式：
1. Markdown 报告（awr_analysis_report.md）
2. Word 报告（awr_analysis_report.docx）

工作流程：解析 AWR → 规则引擎分析 → LLM 分析 → 生成 Markdown → 生成 Word
"""

import json
import re
from datetime import datetime
from pathlib import Path

from .analyzer import write_awr_rule_findings
from .chart_generator import generate_all_charts
from .rules_guide import render_rules_guide_markdown
from .config import (
    AWR_ANALYSIS_DOCX,
    AWR_ANALYSIS_MD,
    AWR_RULE_FINDINGS_MD,
    AWR_SUMMARY_JSON,
    AWR_SUMMARY_MD,
    CHART_DIR,
    CHINESE_NUMBERS,
    COLOR_ACCENT,
    COLOR_ACCENT_LIGHT,
    COLOR_ADVICE_BG,
    COLOR_ADVICE_ACCENT,
    COLOR_BG_DARK,
    COLOR_BG_LIGHT,
    COLOR_BODY,
    COLOR_BORDER,
    COLOR_CONCLUSION_BG,
    COLOR_CONCLUSION_ACCENT,
    COLOR_EVIDENCE_BG,
    COLOR_EVIDENCE_ACCENT,
    COLOR_H2,
    COLOR_H3,
    COLOR_MUTED,
    COLOR_PRIMARY,
    COLOR_RISK_BG,
    COLOR_RISK_ACCENT,
    COLOR_RISK_HIGH,
    COLOR_RISK_LOW,
    COLOR_RISK_MEDIUM,
    COLOR_RISK_NONE,
    COLOR_TABLE_HEAD,
    COLOR_TABLE_ALT,
    COLOR_TITLE,
    FONT_CN,
    FONT_MONO,
    OUTPUT_DIR,
)
from .llm_client import ask_llm, build_awr_prompt, discover_local_models, preferred_model
from .parser import write_awr_summary


# ══════════════════════════════════════════════
# Markdown 报告（直接保存 LLM 输出）
# ══════════════════════════════════════════════


def write_markdown_report(markdown_content: str) -> Path:
    """将 LLM 生成的 Markdown 写入文件。"""
    AWR_ANALYSIS_MD.write_text(markdown_content, encoding="utf-8")
    return AWR_ANALYSIS_MD


# ══════════════════════════════════════════════
# Word 报告（python-docx）— 底层工具函数
# ══════════════════════════════════════════════


def rgb(hex_color: str):
    """字符串色值 → docx RGBColor。"""
    from docx.shared import RGBColor

    value = hex_color.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, size_pt=10.5, bold=False, color=COLOR_BODY, font_name=FONT_CN):
    """设置 run 字体属性。"""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = rgb(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)


def shade_cell(cell, color: str):
    """设置单元格底色。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, right=100, bottom=80, left=100):
    """设置单元格边距。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("right", right), ("bottom", bottom), ("left", left)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLOR_BORDER):
    """设置表格边框。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, size_pt=9, bold=False, color=COLOR_BODY, font_name=FONT_CN):
    """设置单元格文本。"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    add_runs_with_bold(paragraph, str(text), size_pt=size_pt, default_bold=bold, color=color, font_name=font_name)


def add_runs_with_bold(paragraph, text: str, size_pt=10.5, default_bold=False, color=COLOR_BODY, font_name=FONT_CN):
    """添加段落文本（支持 **粗体** 标记）。"""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = default_bold
        content = part
        if part.startswith("**") and part.endswith("**"):
            bold = True
            content = part[2:-2]
        run = paragraph.add_run(content)
        set_run_font(run, size_pt=size_pt, bold=bold, color=color, font_name=font_name)


def format_paragraph(paragraph, size_pt=10.5, color=COLOR_BODY):
    """格式化段落间距。"""
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt, bold=bool(run.bold), color=color)


# ══════════════════════════════════════════════
# 专业排版辅助
# ══════════════════════════════════════════════


def add_accent_bar(paragraph, color=COLOR_ACCENT, size="8", width_percent=30):
    """在段落下方添加金色强调线（专业排版装饰）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    p_pr.append(border)

    # Also set space after for visual breathing room
    paragraph.paragraph_format.space_after = Pt(10)


def add_page_header_footer(document, basic_info: dict):
    """添加专业页眉和页脚（公司/报告标题 + 页码）。"""
    from docx.enum.section import WD_HEADER_FOOTER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    section = document.sections[0]

    # ── 页眉：左对齐报告标题 + 右对齐分隔线 ──
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.clear()

    # Add a thin bottom border to header
    p_pr = header_para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_ACCENT)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    db_name = basic_info.get("DB Name") or "Oracle Database"
    run = header_para.add_run(f"Oracle AWR 性能分析报告  |  {db_name}")
    set_run_font(run, size_pt=8.5, bold=False, color=COLOR_MUTED)

    # ── 页脚：居中页码 ──
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = 1  # center

    # Add thin top border to footer
    fp_pr = footer_para._p.get_or_add_pPr()
    fp_bdr = OxmlElement("w:pBdr")
    ftop = OxmlElement("w:top")
    ftop.set(qn("w:val"), "single")
    ftop.set(qn("w:sz"), "4")
    ftop.set(qn("w:space"), "1")
    ftop.set(qn("w:color"), COLOR_BORDER)
    fp_bdr.append(ftop)
    fp_pr.append(fp_bdr)

    # PAGE field
    run = footer_para.add_run("第 ")
    set_run_font(run, size_pt=8.5, color=COLOR_MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    run2 = footer_para.add_run("PAGE")
    set_run_font(run2, size_pt=8.5, color=COLOR_MUTED)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run2._r.append(fld_char_separate)
    run3 = footer_para.add_run("1")
    set_run_font(run3, size_pt=8.5, color=COLOR_MUTED)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_char_end)
    run4 = footer_para.add_run(" 页")
    set_run_font(run4, size_pt=8.5, color=COLOR_MUTED)


# ══════════════════════════════════════════════
# 文档设置
# ══════════════════════════════════════════════


def create_document():
    """创建 Word 文档对象。"""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("当前 Python 环境未安装 python-docx，请先安装 python-docx") from exc
    document = Document()
    setup_document_styles(document)
    return document


def setup_document_styles(document):
    """设置文档全局样式（专业金融报告风格）。"""
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.start_type = WD_SECTION_START.NEW_PAGE

    # Set page background color
    bg = document.element.find(qn("w:background"))
    if bg is None:
        from docx.oxml import OxmlElement
        bg = OxmlElement("w:background")
        document.element.append(bg)
    bg.set(qn("w:color"), "FFFFFF")

    normal = document.styles["Normal"]
    normal.font.name = FONT_CN
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_before = Pt(1)
    normal.paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════
# Markdown 内容抽取辅助
# ══════════════════════════════════════════════


def load_awr_basic_info() -> dict:
    """从 awr_summary.json 加载 AWR 基本信息。"""
    if not AWR_SUMMARY_JSON.exists():
        return {}
    try:
        payload = json.loads(AWR_SUMMARY_JSON.read_text(encoding="utf-8"))
        return payload.get("basic_info") or {}
    except Exception:
        return {}


def split_table_row(line: str) -> list[str]:
    """解析 Markdown 表格行。"""
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def is_table_separator(line: str) -> bool:
    """判断是否为 Markdown 表格分隔行。"""
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def is_table_line(line: str) -> bool:
    """判断是否为 Markdown 表格行。"""
    return line.strip().startswith("|") and line.strip().endswith("|")


def extract_section(markdown: str, title_pattern: str) -> str:
    """从 Markdown 中提取指定章节内容。"""
    pattern = re.compile(
        rf"(?ims)^##\s*(?:\d+[.、]\s*)?{re.escape(title_pattern)}\s*$\n(.*?)(?=^##\s+|\Z)"
    )
    match = pattern.search(markdown)
    return match.group(1).strip() if match else ""


def extract_key_conclusions(markdown: str) -> list[str]:
    """提取总体结论列表。"""
    section = extract_section(markdown, "总体结论")
    source = section or markdown
    conclusions = []
    for line in source.splitlines():
        stripped = line.strip()
        if stripped.startswith(("- ", "* ")):
            conclusions.append(stripped[2:].strip())
        elif re.match(r"^\d+[.、]\s+", stripped):
            conclusions.append(re.sub(r"^\d+[.、]\s+", "", stripped).strip())
        if len(conclusions) >= 5:
            break
    if not conclusions:
        for line in source.splitlines():
            stripped = line.strip()
            if stripped and not stripped.startswith(("#", "|")):
                conclusions.append(stripped)
            if len(conclusions) >= 5:
                break
    return conclusions or ["Markdown 报告未识别到明确总体结论，请结合 AWR 摘要复核。"]


def extract_first_table_after(markdown: str, title_pattern: str) -> list[str]:
    """提取指定章节后的第一个表格。"""
    section = extract_section(markdown, title_pattern)
    lines = section.splitlines()
    for index, line in enumerate(lines):
        if is_table_line(line) and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            for next_line in lines[index + 2 :]:
                if not is_table_line(next_line):
                    break
                table_lines.append(next_line)
            return table_lines
    return []


# ══════════════════════════════════════════════
# Word 组件渲染
# ══════════════════════════════════════════════


def add_cover_page(document, basic_info: dict):
    """添加专业封面页（金融报告风格）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    # ── 顶部金色装饰线 ──
    for _ in range(3):
        document.add_paragraph()

    top_line = document.add_paragraph()
    top_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top_line.add_run("━" * 60)
    set_run_font(run, size_pt=8, color=COLOR_ACCENT)

    document.add_paragraph()

    # ── 主标题 ──
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Oracle AWR 性能分析报告")
    set_run_font(run, size_pt=24, bold=True, color=COLOR_PRIMARY)

    document.add_paragraph()

    # ── 副标题 ──
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于 AWR / ASH / OS 指标的数据库性能诊断与优化建议")
    set_run_font(run, size_pt=12, color=COLOR_MUTED)

    document.add_paragraph()

    # ── 底部金色装饰线 ──
    bottom_line = document.add_paragraph()
    bottom_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = bottom_line.add_run("━" * 60)
    set_run_font(run, size_pt=8, color=COLOR_ACCENT)

    # ── 分隔空间 ──
    for _ in range(4):
        document.add_paragraph()

    # ── 数据库信息表 ──
    info = [
        ("数据库名称", basic_info.get("DB Name") or "详见 AWR 摘要"),
        ("实例名称", basic_info.get("Instance Name") or "详见 AWR 摘要"),
        ("主机名称", basic_info.get("Host Name") or "详见 AWR 摘要"),
        (
            "分析时间窗口",
            f"{basic_info.get('Begin Snap') or '详见 AWR 摘要'} — {basic_info.get('End Snap') or '详见 AWR 摘要'}",
        ),
        ("AI 分析模型", basic_info.get("Model Name") or "DeepSeek"),
        ("报告生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]

    info_table = document.add_table(rows=len(info), cols=2)
    info_table.autofit = True
    set_table_borders(info_table, color="FFFFFF")
    for index, (key, value) in enumerate(info):
        # 标签列（右对齐）
        set_cell_text(info_table.rows[index].cells[0], key, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
        set_cell_margins(info_table.rows[index].cells[0], top=60, bottom=60, left=200, right=60)
        info_table.rows[index].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 值列
        set_cell_text(info_table.rows[index].cells[1], value, size_pt=10.5, color=COLOR_BODY)
        set_cell_margins(info_table.rows[index].cells[1], top=60, bottom=60, left=60, right=200)

    document.add_page_break()


def add_numbered_heading(document, text: str, level: int, force_text: str | None = None):
    """添加带编号的标题（金融报告风格 — 左饰线 + 间距）。"""
    from docx.shared import Pt

    clean = re.sub(r"^\d+[.、]\s*", "", text).strip()
    if level == 1:
        title = force_text or clean
        paragraph = document.add_paragraph()

        # Add left accent bar via paragraph border
        p_pr = paragraph._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), COLOR_ACCENT)
        p_bdr.append(left)
        # Also add bottom line
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), COLOR_BORDER)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

        run = paragraph.add_run(title)
        set_run_font(run, size_pt=16, bold=True, color=COLOR_TITLE)
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.3
        return

    if level == 2:
        size_pt, color = 12, COLOR_H2
    else:
        size_pt, color = 11, COLOR_H3

    paragraph = document.add_paragraph()
    run = paragraph.add_run(clean)
    set_run_font(run, size_pt=size_pt, bold=True, color=color)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)


def contains_emphasis_keyword(text: str) -> bool:
    """判断是否包含强调关键词。"""
    keywords = ("核心结论", "关键发现", "风险", "建议", "现有 AWR 无法证明", "需要补充")
    return any(keyword in text for keyword in keywords)


def is_special_heading(text: str) -> bool:
    return any(keyword in text for keyword in ("核心结论", "风险提示", "后续需要补充", "专家建议"))


def special_title(text: str) -> str:
    if "核心结论" in text:
        return "核心结论"
    if "风险提示" in text or "风险" in text:
        return "风险提示"
    if "后续" in text or "需要补充" in text:
        return "后续需要补充的证据"
    return "专家建议"


def special_block_type(text: str) -> str:
    if "核心结论" in text:
        return "conclusion"
    if "风险" in text:
        return "risk"
    if "后续" in text or "需要补充" in text:
        return "evidence"
    return "advice"


def collect_block(lines: list[str], start: int) -> tuple[list[str], int]:
    block = []
    index = start
    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith("#") or (
            is_table_line(stripped) and index + 1 < len(lines) and is_table_separator(lines[index + 1].strip())
        ):
            break
        if stripped:
            block.append(re.sub(r"^[-*]\s+", "", stripped))
        index += 1
    return block, index


def add_special_block(document, title: str, content, block_type="conclusion"):
    """添加特殊块（结论/风险/证据/建议）— 专业卡片样式。"""
    colors = {
        "conclusion": (COLOR_CONCLUSION_BG, COLOR_CONCLUSION_ACCENT),
        "risk": (COLOR_RISK_BG, COLOR_RISK_ACCENT),
        "evidence": (COLOR_EVIDENCE_BG, COLOR_EVIDENCE_ACCENT),
        "advice": (COLOR_ADVICE_BG, COLOR_ADVICE_ACCENT),
    }
    fill, accent = colors.get(block_type, colors["conclusion"])
    items = content if isinstance(content, list) else [str(content)]
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    set_table_borders(table, color=fill)

    # Left accent stripe
    shade_cell(table.rows[0].cells[0], accent)
    set_cell_margins(table.rows[0].cells[0], left=30, right=30)
    table.rows[0].cells[0].width = 1

    # Content area
    shade_cell(table.rows[0].cells[1], fill)
    set_cell_margins(table.rows[0].cells[1], top=100, right=140, bottom=80, left=140)

    paragraph = table.rows[0].cells[1].paragraphs[0]
    run = paragraph.add_run(title)
    set_run_font(run, size_pt=11, bold=True, color=COLOR_PRIMARY)

    for item in items:
        p = table.rows[0].cells[1].add_paragraph(style="List Bullet" if len(items) > 1 else None)
        add_runs_with_bold(p, item, size_pt=10, default_bold=True, color=COLOR_BODY)
        format_paragraph(p, size_pt=10)

    document.add_paragraph()


def add_markdown_table(document, rows: list[str]):
    """将 Markdown 表格转换为 Word 表格（专业样式）。"""
    if len(rows) < 2:
        return
    headers = split_table_row(rows[0])
    data_rows = [split_table_row(row) for row in rows[2:] if is_table_line(row)]
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = True
    table.alignment = 1  # center

    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, size_pt=9, bold=True, color="FFFFFF")
    for data in data_rows:
        row = table.add_row()
        for index in range(len(headers)):
            set_cell_text(row.cells[index], data[index] if index < len(data) else "", size_pt=9)
    style_table(table, headers)
    document.add_paragraph()


def style_table(table, headers: list[str]):
    """设置表格样式（金融报告风格 — 深色表头、斑马纹、风险色标、圆角边框）。"""
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            set_cell_margins(cell, top=60, right=80, bottom=60, left=80)
            if row_index == 0:
                # Header row — dark background
                shade_cell(cell, COLOR_TABLE_HEAD)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size_pt=9, bold=True, color="FFFFFF")
                continue
            # Zebra striping
            if row_index % 2 == 0:
                shade_cell(cell, COLOR_TABLE_ALT)
            # Risk coloring
            text = cell.text.strip()
            header = headers[cell_index] if cell_index < len(headers) else ""
            if "风险" in header or "等级" in header:
                if "高" in text:
                    shade_cell(cell, COLOR_RISK_BG)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run, color=COLOR_RISK_ACCENT, bold=True)
                elif "中" in text:
                    shade_cell(cell, "FEF5E7")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run, color=COLOR_RISK_MEDIUM, bold=True)
                elif "低" in text:
                    shade_cell(cell, "E8F0E8")
                elif "未" in text or "信息不足" in text:
                    shade_cell(cell, "F2F4F4")


def normalize_heading_text(text: str) -> str:
    """标准化标题文本。"""
    text = re.sub(r"^\d+[.、]\s*", "", text).strip()
    aliases = {"风险等级": "风险等级判断"}
    return aliases.get(text, text)


def insert_chart_if_needed(document, title: str):
    """在对应章节后插入图表（如果存在）。"""
    chart_map = {
        "负载": "awr_load_profile.png",
        "等待": "top_wait_events.png",
        "SQL": "top_sql.png",
        "CPU": "host_cpu.png",
        "命中率": "instance_efficiency.png",
        "内存": "instance_efficiency.png",
        "缓存": "instance_efficiency.png",
        "Segment": "top_segments.png",
        "段": "top_segments.png",
        "热点": "top_segments.png",
    }
    for keyword, filename in chart_map.items():
        if keyword in title:
            chart_path = CHART_DIR / filename
            if chart_path.exists():
                document.add_picture(str(chart_path), width=14_400_000)  # ~15.24cm
                document.add_paragraph()
                break


def add_executive_summary(document, markdown: str):
    """添加执行摘要（封面后第一页）— 结论卡片 + 风险表格。"""
    conclusions = extract_key_conclusions(markdown)
    add_numbered_heading(document, "结论摘要", 1, force_text="结论摘要")
    add_special_block(document, "核心结论", conclusions, block_type="conclusion")

    risk_table = extract_first_table_after(markdown, "风险等级") or extract_first_table_after(
        markdown, "风险等级判断"
    )
    if risk_table:
        add_numbered_heading(document, "风险等级", 2, force_text="风险等级")
        add_markdown_table(document, risk_table)
    else:
        add_special_block(
            document, "风险提示", ["Markdown 报告中未识别到风险等级表，请结合规则引擎结果复核。"], block_type="risk"
        )
    document.add_page_break()


def add_markdown_content(document, markdown: str):
    """将 Markdown 内容逐段渲染到 Word 文档。"""
    lines = markdown.splitlines()
    section_index = 0
    index = 0
    while index < len(lines):
        stripped = lines[index].rstrip().strip()
        if not stripped:
            index += 1
            continue
        if is_table_line(stripped) and index + 1 < len(lines) and is_table_separator(lines[index + 1].strip()):
            table_lines = [stripped, lines[index + 1].strip()]
            index += 2
            while index < len(lines) and is_table_line(lines[index].strip()):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(document, table_lines)
            continue
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            if title != "Oracle AWR 性能分析报告":
                paragraph = document.add_paragraph()
                paragraph.alignment = 1  # center
                run = paragraph.add_run(title)
                set_run_font(run, size_pt=20, bold=True, color=COLOR_PRIMARY)
        elif stripped.startswith("## "):
            section_index += 1
            title = normalize_heading_text(stripped[3:].strip())
            numbered = (
                f"{CHINESE_NUMBERS[section_index - 1] if section_index <= len(CHINESE_NUMBERS) else section_index}、"
                f"{title}"
            )
            add_numbered_heading(document, numbered, 1, force_text=numbered)
            insert_chart_if_needed(document, title)
        elif stripped.startswith("### "):
            add_numbered_heading(document, stripped[4:].strip(), 2)
        elif is_special_heading(stripped):
            block_lines, index = collect_block(lines, index + 1)
            add_special_block(
                document,
                special_title(stripped),
                block_lines,
                block_type=special_block_type(stripped),
            )
            continue
        elif stripped.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            add_runs_with_bold(
                paragraph,
                stripped[2:].strip(),
                size_pt=10,
                default_bold=contains_emphasis_keyword(stripped),
            )
            format_paragraph(paragraph, size_pt=10)
        else:
            paragraph = document.add_paragraph()
            add_runs_with_bold(
                paragraph,
                stripped,
                size_pt=10,
                default_bold=contains_emphasis_keyword(stripped),
            )
            format_paragraph(paragraph, size_pt=10)
        index += 1


# ══════════════════════════════════════════════
# 主入口
# ══════════════════════════════════════════════


def markdown_to_word(markdown_path: Path | None = None, model_name: str | None = None) -> dict:
    """将 AWR 分析 Markdown 报告转为 Word 文档。

    Args:
        markdown_path: Markdown 文件路径，None 则使用默认路径
        model_name: AI 分析模型名称，None 则从 basic_info 读取

    Returns:
        包含 word_path 等信息的字典
    """
    md_path = markdown_path or AWR_ANALYSIS_MD
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown 报告不存在：{md_path}")

    markdown = md_path.read_text(encoding="utf-8")
    basic_info = load_awr_basic_info()
    if model_name:
        basic_info["Model Name"] = model_name

    document = create_document()
    add_page_header_footer(document, basic_info)
    add_cover_page(document, basic_info)
    add_executive_summary(document, markdown)
    add_markdown_content(document, markdown)

    # ── 附录：规则引擎说明 ──
    document.add_page_break()
    rules_guide_md = render_rules_guide_markdown()
    add_markdown_content(document, rules_guide_md)

    document.save(str(AWR_ANALYSIS_DOCX))
    return {
        "word_path": str(AWR_ANALYSIS_DOCX),
        "summary": basic_info,
    }


# ══════════════════════════════════════════════
# 完整分析流程
# ══════════════════════════════════════════════


def run_full_analysis(
    awr_path: str,
    url: str | None = None,
    model: str | None = None,
    api_key: str = "",
    skip_llm: bool = False,
) -> dict:
    """执行完整的 AWR 分析流水线。

    流程：
        1. 解析 AWR HTML 报告 → 结构化摘要
        2. 规则引擎分析 → 发现 & 指标
        3. 调用 DeepSeek → 报告 Markdown
        4. 保存 Markdown 报告
        5. 生成 Word 报告

    Args:
        awr_path: AWR HTML 报告路径
        url: DeepSeek API 地址
        model: 模型名称
        skip_llm: 是否跳过 LLM 分析（仅执行解析 + 规则引擎）

    Returns:
        包含所有步骤结果的字典
    """
    path = Path(awr_path).expanduser().resolve()
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"AWR 文件不存在：{path}")

    resolved_url = url or DEFAULT_DEEPSEEK_URL
    resolved_model = model or preferred_model(resolved_url)

    # 1. 解析 AWR
    print(f"📖 解析 AWR 报告：{path.name}")
    summary = write_awr_summary(path)

    # 1b. 生成图表（基于结构化摘要）
    print("📊 生成性能可视化图表...")
    try:
        charts = generate_all_charts(summary)
    except Exception as exc:
        print(f"   ⚠️ 图表生成异常（可忽略）：{exc}")
        charts = {}

    # 2. 规则引擎分析
    print("🔍 执行规则引擎分析...")
    rule_findings = write_awr_rule_findings(summary)

    result = {
        "source_file": str(path),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "summary": summary,
        "rule_findings": rule_findings,
        "llm_analysis": None,
        "markdown_path": str(AWR_ANALYSIS_MD),
        "word_path": str(AWR_ANALYSIS_DOCX),
    }

    if skip_llm:
        print("⏭️ 跳过 LLM 分析")
        return result

    # 3. LLM 分析
    print(f"🤖 调用本地模型（{resolved_model}）分析...")
    summary_md = AWR_SUMMARY_MD.read_text(encoding="utf-8")
    rule_findings_md = AWR_RULE_FINDINGS_MD.read_text(encoding="utf-8")
    rules_guide_md = render_rules_guide_markdown()
    prompt = build_awr_prompt(summary_md, rule_findings_md, rules_guide_md)
    answer = ask_llm(resolved_url, resolved_model, prompt, api_key=api_key)

    # 4. 保存 Markdown 报告
    print("📝 保存 Markdown 报告...")
    write_markdown_report(answer)

    # 5. 生成 Word 报告
    print("📄 生成 Word 报告...")
    try:
        word_result = markdown_to_word(model_name=resolved_model)
        result["word_path"] = word_result["word_path"]
    except Exception as exc:
        print(f"⚠️ Word 报告生成失败（可忽略）：{exc}")
        result["word_path"] = None

    result["llm_analysis"] = {
        "url": resolved_url,
        "model": resolved_model,
        "answer": answer,
    }

    print(f"✅ 分析完成！")
    print(f"   Markdown 报告：{AWR_ANALYSIS_MD}")
    print(f"   Word 报告：{AWR_ANALYSIS_DOCX}")
    return result
